import streamlit as st
import io
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==========================================
# 核心功能 1：软著代码清洗与 60 页排版引擎
# ==========================================
def generate_copyright_word(files, software_name, version):
    # 1. 读取并清洗代码（去除完全空白的行）
    all_lines = []
    for f in files:
        text = f.read().decode('utf-8', errors='ignore')
        for line in text.split('\n'):
            if line.strip():  # 只要不是纯纯的空行，就保留（保留原始缩进）
                all_lines.append(line.rstrip())
    
    total_raw_lines = len(all_lines)
    
    # 2. 截取前 1500 行和后 1500 行 (刚好 60 页，每页 50 行)
    if total_raw_lines > 3000:
        selected_lines = all_lines[:1500] + all_lines[-1500:]
    else:
        selected_lines = all_lines
        
    # 3. 构建 Word 文档
    doc = Document()
    
    # 设置 A4 纸张与边距 (版权局标准边距)
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    
    # 设置页眉 (软件名称 + 版本号，居中)
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{software_name} {version}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.style.font.name = '黑体'
    header_para.style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 设置全局默认字体 (五号字体，宋体/Courier New)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New' # 代码用等宽字体最好看
    font.size = Pt(10.5)      # 五号字
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 4. 精准按页写入代码 (每 50 行强制分页)
    current_para = doc.add_paragraph()
    # 强行锁定行距，防止 Word 自动调整导致一页超行
    current_para.paragraph_format.line_spacing = Pt(12) 
    current_para.paragraph_format.space_after = Pt(0)
    current_para.paragraph_format.space_before = Pt(0)
    
    lines_in_page = 0
    for i, line in enumerate(selected_lines):
        current_para.add_run(line + '\n')
        lines_in_page += 1
        
        # 满 50 行，且不是最后一行时，强制打断换页！
        if lines_in_page == 50 and i != len(selected_lines) - 1:
            doc.add_page_break()
            current_para = doc.add_paragraph()
            current_para.paragraph_format.line_spacing = Pt(12)
            current_para.paragraph_format.space_after = Pt(0)
            current_para.paragraph_format.space_before = Pt(0)
            lines_in_page = 0
            
    # 输出到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output, total_raw_lines, len(selected_lines)

# ==========================================
# 核心功能 2：专利交底书框架生成引擎
# ==========================================
def generate_patent_disclosure(title, background, problem, solution, effects):
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12) # 小四
    
    doc.add_heading('发明专利技术交底书', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('一、 发明名称', level=1)
    doc.add_paragraph(title)
    
    doc.add_heading('二、 背景技术 (现有技术的缺点)', level=1)
    doc.add_paragraph(background)
    
    doc.add_heading('三、 本发明要解决的技术问题', level=1)
    doc.add_paragraph(problem)
    
    doc.add_heading('四、 技术方案 (核心架构与步骤)', level=1)
    doc.add_paragraph(solution)
    
    doc.add_heading('五、 有益效果', level=1)
    doc.add_paragraph(effects)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# Streamlit UI 界面
# ==========================================
st.set_page_config(page_title="IP Auto-Framer | 软著专利神器", page_icon="📜", layout="wide")

st.title("📜 IP Auto-Framer 知识产权排版神器")
st.markdown("专为中国版权保护中心 (CPCC) 和国家知识产权局量身定制的无脑格式化工具。")

tab1, tab2 = st.tabs(["💻 软著代码 60 页全自动清洗", "📝 发明专利交底书生成器"])

# ------------- 软著模块 -------------
with tab1:
    st.markdown("### 🛠️ 软著源代码清洗与排版")
    st.info("💡 规则提示：版权局要求提交前 30 页和后 30 页，每页严格 50 行，且不能有大量空行。本工具将自动帮您完成这一切！")
    
    col1, col2 = st.columns(2)
    soft_name = col1.text_input("软件全称 (用于生成页眉)", value="基于大模型的全自动生信分析系统")
    soft_version = col2.text_input("版本号", value="V1.0")
    
    uploaded_codes = st.file_uploader("上传您的代码文件 (支持多个 .py, .java, .cpp 等后缀)", accept_multiple_files=True)
    
    if st.button("🚀 一键生成 60 页软著代码 Word", type="primary"):
        if not uploaded_codes:
            st.warning("请先上传代码文件！")
        else:
            with st.spinner("正在疯狂清洗空行并强行分页..."):
                word_io, total_lines, final_lines = generate_copyright_word(uploaded_codes, soft_name, soft_version)
                
                st.success(f"✅ 处理完成！原始总代码行数：{total_lines} 行。清洗后提交流水线总行数：{final_lines} 行。")
                
                st.download_button(
                    label="📥 点击下载软著源代码.docx",
                    data=word_io,
                    file_name=f"{soft_name}_软著源代码.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ------------- 专利模块 -------------
with tab2:
    st.markdown("### 📝 标准发明专利交底书生成器")
    st.markdown("不用对着白纸发呆，按照填空题写出您的 Idea，一键生成代理人最爱的标准格式。")
    
    pat_title = st.text_input("1. 拟定专利名称", placeholder="例如：一种基于大语言模型的质谱仪智能控制系统及方法")
    pat_bg = st.text_area("2. 背景技术（目前别人是怎么做的？有什么痛点？）", height=100)
    pat_prob = st.text_area("3. 核心技术问题（您想解决什么致命痛点？）", height=80)
    pat_sol = st.text_area("4. 技术方案（您的核心三步走架构是什么？怎么实现的？）", height=200)
    pat_eff = st.text_area("5. 有益效果（带来了什么降本增效的奇迹？）", height=100)
    
    if st.button("🚀 生成专利交底书 Word 模板"):
        if not pat_title:
            st.warning("至少把标题填上吧！")
        else:
            pat_io = generate_patent_disclosure(pat_title, pat_bg, pat_prob, pat_sol, pat_eff)
            st.success("✅ 交底书框架生成成功！")
            st.download_button(
                label="📥 点击下载专利交底书.docx",
                data=pat_io,
                file_name=f"{pat_title}_技术交底书.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
